# Python
import sys
import sqlite3
import os
import re
from typing import List, Tuple
from PyQt6.QtWidgets import (
    QApplication,
    QWidget,
    QLabel,
    QPushButton,
    QVBoxLayout,
    QMessageBox,
    QMainWindow,
    QStatusBar,
    QMenuBar,
    QFileDialog,
    QTextEdit,
    QHBoxLayout,
    QProgressBar,
    QListWidget,
    QDialog,
    QDialogButtonBox,
    QListWidgetItem,
    QInputDialog,
    )
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QObject
from PyQt6.QtGui import QDragEnterEvent, QDropEvent
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
from cryptography.hazmat.primitives.hashes import SHA256
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC

import pdfplumber.utils
import docx.opc.exceptions

# Document parsing libraries
import pdfplumber
from docx import Document  # python-docx
import pytesseract
from PIL import Image  # Pillow for image processing with Tesseract
import pandas as pd  # Pandas for Excel and CSV

# NLP libraries
import spacy
from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline

# --- Configuration ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATABASE_PATH = os.path.join(BASE_DIR, '..', 'data', 'compliance.db')
HASH_ALGORITHM = SHA256()
ITERATIONS = 100000
OFFLINE_ONLY = True

# --- Tesseract OCR Path (Optional, Windows offline) ---
tess_env = os.environ.get("TESSERACT_EXE")
if tess_env and os.path.isfile(tess_env):
    pytesseract.pytesseract.tesseract_cmd = tess_env

# --- SpaCy Model Loading ---
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    if OFFLINE_ONLY:
        print("SpaCy model 'en_core_web_sm' not found and offline mode is enabled. Install it in this venv before use.")
        nlp = None
    else:
        print("Downloading SpaCy model 'en_core_web_sm'...")
        import spacy.cli
        spacy.cli.download("en_core_web_sm")
        nlp = spacy.load("en_core_web_sm")

# --- Clinical/Biomedical NER Model Loading (offline-first with fallbacks) ---
CLINICAL_NER_MODEL_CANDIDATES = [
    "microsoft/BiomedVLP-CXR-BERT-specialized",
    "d4data/biomedical-ner-all",
    "kamalkraj/BioBERT-NER",
    "dslim/bert-base-NER",
]
clinical_ner_pipeline = None
_loaded_model_name = None
_last_model_error = None
for model_name in CLINICAL_NER_MODEL_CANDIDATES:
    try:
        ner_tokenizer = AutoTokenizer.from_pretrained(model_name, local_files_only=OFFLINE_ONLY)
        ner_model = AutoModelForTokenClassification.from_pretrained(model_name, local_files_only=OFFLINE_ONLY)
        clinical_ner_pipeline = pipeline("ner", model=ner_model, tokenizer=ner_tokenizer, aggregation_strategy="simple")
        _loaded_model_name = model_name
        break
    except Exception as e:
        _last_model_error = e
        continue
if clinical_ner_pipeline is None:
    print("Failed to load a clinical/biomedical NER model in offline mode.")
    if _last_model_error:
        print(f"Last error: {_last_model_error}")
    print("Seed models into the local cache in this virtualenv on an internet-enabled machine, then copy to deployment.")

# --- PHI Scrubber (basic, extendable) ---
def scrub_phi(text: str) -> str:
    if not isinstance(text, str):
        return text
    patterns = [
        (r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}', '[EMAIL]'),
        (r'(\+?\d{1,2}[\s\-.]?)?(\(?\d{3}\)?[ \-.]?\d{3}[\-.]?\d{4})', '[PHONE]'),
        (r'\b\d{3}-\d{2}-\d{4}\b', '[SSN]'),
        (r'\bMRN[:\s]*[A-Za-z0-9\-]{4,}\b', '[MRN]'),
        (r'\b(19|20)\d{2}-/ (0?[1-9]|1[0-2])-/ (0?[1-g]|[12]\d|3[01])\b', '[DATE]'),
        (r'\b(Name|Patient|DOB|Address)[:\s]+[^\n]+', r'\1: [REDACTED]'),
    ]
    out = text
    for pat, repl in patterns:
        out = re.sub(pat, repl, out)
    return out

# --- Helpers: chunking for long texts ---
def chunk_text(text: str, max_chars: int = 4000):
    chunks = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + max_chars, n)
        newline_pos = text.rfind("\n", start, end)
        if newline_pos != -1 and newline_pos > start + 1000:
            end = newline_pos
        chunks.append(text[start:end])
        start = end
    return chunks

# --- Helpers: Standalone Document Parser for Rubrics ---
def parse_document_content(file_path: str) -> List[Tuple[str, str]]:
    if not nlp:
        return [("Error: SpaCy model not loaded. Cannot sentence-split.", "System")]
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        sentences_with_source = []
        if file_extension == '.pdf':
            try:
                with pdfplumber.open(file_path) as pdf:
                    for i, page in enumerate(pdf.pages, start=1):
                        page_text = page.extract_text() or ""
                        doc = nlp(page_text)
                        for sent in doc.sents:
                            if sent.text.strip():
                                sentences_with_source.append((sent.text.strip(), f"Page {i}"))
            except pdfplumber.utils.PDFSyntaxError as e:
                return [(f"Error: Invalid PDF file: {e}", "PDF Parser")]
        elif file_extension == '.docx':
            try:
                doc = Document(file_path)
                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        para_doc = nlp(para.text)
                        for sent in para_doc.sents:
                            if sent.text.strip():
                                sentences_with_source.append((sent.text.strip(), f"Paragraph {i+1}"))
            except docx.opc.exceptions.PackageNotFoundError as e:
                return [(f"Error: Invalid DOCX file: {e}", "DOCX Parser")]
        elif file_extension in ['.xlsx', '.xls', '.csv']:
            text_content = ""
            source_name = "Excel Sheet" if file_extension.startswith('.xls') else "CSV File"
            try:
                if file_extension in ['.xlsx', '.xls']:
                    df = pd.read_excel(file_path)
                    text_content = df.to_string(index=False)
                else:
                    df = pd.read_csv(file_path)
                    text_content = df.to_string(index=False)
                doc = nlp(text_content)
                for sent in doc.sents:
                    if sent.text.strip():
                        sentences_with_source.append((sent.text.strip(), source_name))
            except Exception as e:
                return [(f"Error: Failed to read tabular file: {e}", "Data Parser")]
        elif file_extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']:
            try:
                img = Image.open(file_path)
                if img.mode not in ("RGB", "L"):
                    img = img.convert("RGB")
                ocr_text = pytesseract.image_to_string(img)
                doc = nlp(ocr_text)
                for sent in doc.sents:
                    if sent.text.strip():
                        sentences_with_source.append((sent.text.strip(), "Source: Image (OCR)"))
            except Image.UnidentifiedImageError as e:
                return [(f"Error: Unidentified image file: {e}", "OCR Parser")]
            except Exception as e:
                return [(f"Error: Failed to process image with Tesseract: {e}", "OCR Parser")]
        else:
            return [(f"Error: Unsupported file type: {file_extension}", "File Handler")]
        if not sentences_with_source:
            return [("Info: No text could be extracted from the document.", "System")]
        return sentences_with_source
    except FileNotFoundError:
        return [(f"Error: File not found at {file_path}", "File System")]
    except Exception as e:
        return [(f"Error: An unexpected error occurred: {e}", "System")]

# --- Helpers: Database Initialization ---
def initialize_database():
    try:
        os.makedirs(os.path.dirname(os.path.abspath(DATABASE_PATH)), exist_ok=True)
        with sqlite3.connect(DATABASE_PATH) as conn:
            cur = conn.cursor()
            cur.execute("CREATE TABLE IF NOT EXISTS users (username TEXT PRIMARY KEY, password_hash BLOB NOT NULL, salt BLOB NOT NULL)")
            cur.execute("CREATE TABLE IF NOT EXISTS rubrics (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT NOT NULL UNIQUE, content TEXT NOT NULL, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
            cur.execute("SELECT COUNT(*) FROM rubrics")
            if cur.fetchone()[0] == 0:
                default_rubric_name = "Default Best Practices"
                default_rubric_content = """# General Documentation Best Practices
- All entries must be dated and signed.
- Patient identification must be clear on every page.
- Use of approved abbreviations only.
- Document skilled intervention, not just patient performance.
- Goals must be measurable and time-bound."""
                cur.execute("INSERT INTO rubrics (name, content) VALUES(?, ?)", (default_rubric_name, default_rubric_content))
            username = "test"
            password = "test123"
            salt = os.urandom(16)
            kdf = PBKDF2HMAC(algorithm=HASH_ALGORITHM, length=32, salt=salt, iterations=ITERATIONS)
            password_hash = kdf.derive(password.encode())
            cur.execute("INSERT INTO users (username, password_hash, salt) VALUES (?, ?, ?) ON CONFLICT(username) DO UPDATE SET password_hash=excluded.password_hash, salt=excluded.salt", (username, password_hash, salt))
            conn.commit()
    except sqlite3.Error as e:
        print(f"Failed to initialize database: {e}")

# --- Background Workers ---
class DocumentWorker(QObject):
    finished = pyqtSignal(list)
    error = pyqtSignal(str)
    progress = pyqtSignal(int)

    def __init__(self, file_path: str):
        super().__init__()
        self.file_path = file_path
        self._cancel = False

    def cancel(self):
        self._cancel = True

    def run(self):
        if not nlp:
            self.error.emit("SpaCy model not loaded. Cannot process document.")
            return
        try:
            file_extension = os.path.splitext(self.file_path)[1].lower()
            sentences_with_source = []
            if file_extension == '.pdf':
                with pdfplumber.open(self.file_path) as pdf:
                    total = max(len(pdf.pages), 1)
                    for i, page in enumerate(pdf.pages, start=1):
                        if self._cancel:
                            self.error.emit("Analysis canceled by user.")
                            return
                        page_text = page.extract_text() or ""
                        doc = nlp(page_text)
                        for sent in doc.sents:
                            if sent.text.strip():
                                sentences_with_source.append((sent.text.strip(), f"Page {i}"))
                        self.progress.emit(int((i / total) * 100))
            elif file_extension == '.docx':
                doc = Document(self.file_path)
                total = max(len(doc.paragraphs), 1)
                for i, para in enumerate(doc.paragraphs, start=1):
                    if self._cancel:
                        self.error.emit("Analysis canceled by user.")
                        return
                    if para.text.strip():
                        para_doc = nlp(para.text)
                        for sent in para_doc.sents:
                            if sent.text.strip():
                                sentences_with_source.append((sent.text.strip(), f"Paragraph {i}"))
                    self.progress.emit(int((i / total) * 100))
            else:
                sentences_with_source = parse_document_content(self.file_path)
                if sentences_with_source and sentences_with_source[0][0].startswith("Error:"):
                    self.error.emit(f"{sentences_with_source[0][0]} (Source: {sentences_with_source[0][1]})")
                    return
                self.progress.emit(100)
            if not sentences_with_source:
                self.error.emit('Info: No text could be extracted from the document.')
                return
            self.finished.emit(sentences_with_source)
        except Exception as e:
            import traceback
            self.error.emit(f"Error processing file: {e}\n{traceback.format_exc()}")

class NERWorker(QObject):
    finished = pyqtSignal(str, str)
    error = pyqtSignal(str)
    progress = pyqtSignal(int)

    def __init__(self, text: str):
        super().__init__()
        self.text = text
        self._cancel = False

    def cancel(self):
        self._cancel = True

    def run(self):
        if not clinical_ner_pipeline:
            self.error.emit("Clinical/Biomedical NER model not loaded (offline). Seed models locally and restart.")
            return
        try:
            chunks = chunk_text(self.text)
            aggregated = []
            total = len(chunks) if chunks else 1
            for idx, chunk in enumerate(chunks, start=1):
                if self._cancel:
                    self.error.emit("Analysis canceled by user.")
                    return
                results = clinical_ner_pipeline(chunk)
                for entity in results:
                    aggregated.append(f"Entity: {entity.get('word', '')}| Type: {entity.get('entity_group', '')} | Score: {entity.get('score', 0.0): .2f}")
                self.progress.emit(int((idx / total) * 100))
            output = "\n".join(aggregated) if aggregated else "No entities detected."
            model_name = _loaded_model_name or "Unknown NER model"
            self.finished.emit(model_name, output)
        except Exception as e:
            self.error.emit(f"Error during Clinical/Biomedical NER: {e}")

class AnalysisWorker(QObject):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)
    progress = pyqtSignal(int)

    def __init__(self, doc_sentences_with_source: List[Tuple[str, str]], rubric_content: str, offline_only: bool):
        super().__init__()
        self.doc_sentences_with_source = doc_sentences_with_source
        self.rubric_content = rubric_content
        self.offline_only = offline_only
        self._cancel = False

    def cancel(self):
        self._cancel = True

    def run(self):
        try:
            if self._cancel:
                self.error.emit("Analysis canceled by user.")
                return
            self.progress.emit(10)
            analyzer = SemanticAnalyzer(offline_only=self.offline_only)
            if not analyzer.model:
                self.error.emit("Failed to load semantic analysis model. Check logs or try reinstalling sentence-transformers.")
                return
            if self._cancel:
                return
            self.progress.emit(30)
            results = analyzer.analyze(self.doc_sentences_with_source, self.rubric_content)
            if self._cancel:
                return
            self.progress.emit(90)
            report_lines = ["--- Semantic Rubric Analysis Report ---", ""]
            for res in results:
                status = res['status']
                rule = res['rule']
                score = res['score']
                match = res['match']
                source = res['source']
                report_lines.append(f"[{status}] - Rule: {rule}")
                if status == "MET":
                    report_lines.append(f"     Match (Score: {score:.2f}): {match} (Source: {source})")
                report_lines.append("")
            final_report = "\n".join(report_lines)
            self.progress.emit(100)
            self.finished.emit(final_report)
        except Exception as e:
            import traceback
            self.error.emit(f"An error occurred during rubric analysis: {e}\n{traceback.format_exc()}")

class SemanticAnalyzer:
    def __init__(self, model_name: str = 'all-MiniLM-L6-v2', offline_only: bool = True):
        try:
            from sentence_transformers import SentenceTransformer
            self.model = SentenceTransformer(model_name, local_files_only=offline_only)
        except ImportError:
            self.model = None
        except Exception as e:
            print(f"Failed to load SentenceTransformer model: {e}")
            self.model = None

    def analyze(self, doc_sentences_with_source: List[Tuple[str, str]], rubric_content: str, similarity_threshold: float = 0.5) -> list[dict]:
        if not self.model:
            return [{"rule": "Error", "status": "Semantic model not loaded.", "match": "", "score": 0, "source": ""}]
        from sentence_transformers.util import semantic_search
        rules = [rule.strip() for rule in rubric_content.split('\n') if rule.strip() and not rule.strip().startswith("#")]
        if not doc_sentences_with_source:
            return []
        doc_sentences, doc_sources = zip(*doc_sentences_with_source)
        doc_sentences = list(doc_sentences)
        if not rules or not doc_sentences:
            return []
        rule_embeddings = self.model.encode(rules, convert_to_tensor=True, show_progress_bar=False)
        sentence_embeddings = self.model.encode(doc_sentences, convert_to_tensor=True, show_progress_bar=False)
        hits = semantic_search(rule_embeddings, sentence_embeddings, top_k=1)
        analysis_report = []
        for i, rule_hits in enumerate(hits):
            rule = rules[i]
            result = {"rule": rule, "status": "NOT MET", "match": "No similar sentence found.", "score": 0, "source": ""}
            if rule_hits:
                top_hit = rule_hits[0]
                score = top_hit['score']
                if score >= similarity_threshold:
                    corpus_id = top_hit['corpus_id']
                    result["status"] = "MET"
                    result["match"] = doc_sentences[corpus_id]
                    result["source"] = doc_sources[corpus_id]
                    result["score"] = score
            analysis_report.append(result)
        return analysis_report

class AddRubricSourceDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Select Rubric Source")
        self.source = None
        layout = QVBoxLayout(self)
        layout.addWidget(QLabel("Where would you like to add a rubric from?"))
        self.library_button = QPushButton("From Pre-loaded Library")
        self.library_button.clicked.connect(self.select_library)
        layout.addWidget(self.library_button)
        self.file_button = QPushButton("From Local File")
        self.file_button.clicked.connect(self.select_file)
        layout.addWidget(self.file_button)
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Cancel)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

    def select_library(self):
        self.source = 'library'
        self.accept()

    def select_file(self):
        self.source = 'file'
        self.accept()

class LibrarySelectionDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Select from Library")
        self.selected_path = None
        self.selected_name = None
        layout = QVBoxLayout(self)
        layout.addWidget(QLabel("Select a pre-loaded rubric to add:"))
        self.library_list = QListWidget()
        layout.addWidget(self.library_list)
        self.populate_library_list()
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(self.confirm_selection)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

    def populate_library_list(self):
        rubrics_dir = os.path.join(BASE_DIR, '..', 'resources', 'rubrics')
        if not os.path.isdir(rubrics_dir):
            self.library_list.addItem("No library found.")
            self.library_list.setEnabled(False)
            return
        for filename in os.listdir(rubrics_dir):
            if filename.endswith(".txt"):
                display_name = os.path.splitext(filename)[0].replace('_', ' ').title()
                item = QListWidgetItem(display_name)
                item.setData(Qt.ItemDataRole.UserRole, os.path.join(rubrics_dir, filename))
                self.library_list.addItem(item)

    def confirm_selection(self):
        selected_items = self.library_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Selection Required", "Please select a rubric from the list.")
            return
        item = selected_items[0]
        self.selected_name = item.text()
        self.selected_path = item.data(Qt.ItemDataRole.UserRole)
        self.accept()

class RubricManagerDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Rubric Manager")
        self.setGeometry(150, 150, 400, 300)
        layout = QVBoxLayout(self)
        self.rubric_list = QListWidget()
        layout.addWidget(self.rubric_list)
        self.load_rubrics()
        button_box = QDialogButtonBox()
        self.add_button = button_box.addButton("Add...", QDialogButtonBox.ButtonRole.ActionRole)
        self.remove_button = button_box.addButton("Remove", QDialogButtonBox.ButtonRole.ActionRole)
        close_button = button_box.addButton(QDialogButtonBox.StandardButton.Close)
        close_button.clicked.connect(self.accept)
        self.add_button.clicked.connect(self.add_rubric)
        self.remove_button.clicked.connect(self.remove_rubric)
        layout.addWidget(button_box)

    def load_rubrics(self):
        self.rubric_list.clear()
        try:
            if not os.path.exists(DATABASE_PATH):
                return
            with sqlite3.connect(DATABASE_PATH) as conn:
                cur = conn.cursor()
                cur.execute("SELECT id, name FROM rubrics ORDER BY name ASC")
                for rubric_id, name in cur.fetchall():
                    item = QListWidgetItem(name)
                    item.setData(Qt.ItemDataRole.UserRole, rubric_id)
                    self.rubric_list.addItem(item)
        except sqlite3.Error as e:
            QMessageBox.critical(self, "Error", f"Failed to load rubrics from database:\n{e}")

    def add_rubric(self):
        source_dialog = AddRubricSourceDialog(self)
        if not source_dialog.exec():
            return
        if source_dialog.source == 'file':
            self.add_rubric_from_file()
        elif source_dialog.source == 'library':
            self.add_rubric_from_library()

    def add_rubric_from_file(self):
        rubric_name, ok = QInputDialog.getText(self, "Add Rubric From File", "Enter a unique name for the new rubric:")
        if not (ok and rubric_name):
            return
        file_path, _ = QFileDialog.getOpenFileName(self, 'Select Rubric Document', '', 'Supported Files(*.pdf *.docx *.xlsx *.xls *.csv *.png *.jpg *.jpeg *.gif *.bmp *.tiff);;All Files(*.*)')
        if not file_path:
            return
        content = parse_document_content(file_path)
        if content[0][0].startswith("Error:"):
            QMessageBox.critical(self, "Error", f"Failed to parse rubric document:\n{content[0][0]}")
            return
        content_str = "\n".join([text for text, source in content])
        try:
            with sqlite3.connect(DATABASE_PATH) as conn:
                cur = conn.cursor()
                cur.execute("INSERT INTO rubrics (name, content) VALUES(?, ?)", (rubric_name, content_str))
                conn.commit()
                QMessageBox.information(self, "Success", f"Rubric '{rubric_name}' added successfully.")
                self.load_rubrics()
        except sqlite3.IntegrityError:
            QMessageBox.critical(self, "Error", f"A rubric with the name '{rubric_name}' already exists. Please choose a unique name.")
        except sqlite3.Error as e:
            QMessageBox.critical(self, "Database Error", f"Failed to save rubric to database:\n{e}")

    def add_rubric_from_library(self):
        lib_dialog = LibrarySelectionDialog(self)
        if not lib_dialog.exec():
            return
        rubric_name = lib_dialog.selected_name
        rubric_path = lib_dialog.selected_path
        content = parse_document_content(rubric_path)
        if content[0][0].startswith("Error:"):
            QMessageBox.critical(self, "Error", f"Failed to parse library rubric:\n{content[0][0]}")
            return
        content_str = "\n".join([text for text, source in content])
        try:
            with sqlite3.connect(DATABASE_PATH) as conn:
                cur = conn.cursor()
                cur.execute("INSERT INTO rubrics (name, content) VALUES(?, ?)", (rubric_name, content_str))
                conn.commit()
                QMessageBox.information(self, "Success", f"Rubric '{rubric_name}' added from library.")
                self.load_rubrics()
        except sqlite3.IntegrityError:
            QMessageBox.warning(self, "Already Exists", f"The library rubric '{rubric_name}' is already in your database.")
        except sqlite3.Error as e:
            QMessageBox.critical(self, "Database Error", f"Failed to save rubric to database:\n{e}")

    def remove_rubric(self):
        selected_items = self.rubric_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Remove Rubric", "Please select a rubric to remove.")
            return
        item = selected_items[0]
        rubric_id = item.data(Qt.ItemDataRole.UserRole)
        rubric_name = item.text()
        reply = QMessageBox.question(self, "Confirm Deletion", f"Are you sure you want to permanently delete the rubric '{rubric_name}'?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            try:
                with sqlite3.connect(DATABASE_PATH) as conn:
                    cur = conn.cursor()
                    cur.execute("DELETE FROM rubrics WHERE id = ?", (rubric_id,))
                    conn.commit()
                QMessageBox.information(self, "Success", f"Rubric '{rubric_name}' has been deleted.")
                self.load_rubrics()
            except Exception as e:
                QMessageBox.critical(self, "Database Error", f"Failed to delete rubric:\n{e}")

class MainApplicationWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Therapy Compliance Analyzer')
        self.setGeometry(100, 100, 1024, 768)
        self.scrub_before_display = True
        self._current_raw_text = ""
        self._current_sentences_with_source: List[Tuple[str, str]] = []
        self._current_entities_spacy = ""
        self._current_entities_transformer = ""
        self.setAcceptDrops(True)
        self.menu_bar = QMenuBar(self)
        self.setMenuBar(self.menu_bar)
        self.file_menu = self.menu_bar.addMenu('File')
        self.file_menu.addAction('Exit', self.close)
        self.tools_menu = self.menu_bar.addMenu('Tools')
        self.tools_menu.addAction('Initialize Database', initialize_database)
        self.tools_menu.addAction('Quickstart', self.show_quickstart)
        self.tools_menu.addAction('Verify Offline Readiness', self.verify_offline_readiness)
        self.admin_menu = self.menu_bar.addMenu('Admin Options')
        self.toggle_scrub_action = self.admin_menu.addAction('Scrub PHI before display (recommended)')
        self.toggle_scrub_action.setCheckable(True)
        self.toggle_scrub_action.setChecked(self.scrub_before_display)
        self.toggle_scrub_action.toggled.connect(self._toggle_scrub_setting)
        self.help_menu = self.menu_bar.addMenu('Help')
        self.help_menu.addAction('Show Paths', self.show_paths)
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage('Ready')
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        main_layout = QVBoxLayout(self.central_widget)
        button_layout = QHBoxLayout()
        self.upload_button = QPushButton('Upload Document')
        self.upload_button.clicked.connect(self.open_file_dialog)
        button_layout.addWidget(self.upload_button)
        self.clear_button = QPushButton('Clear Display')
        self.clear_button.clicked.connect(self.clear_document_display)
        button_layout.addWidget(self.clear_button)
        self.generate_pdf_button = QPushButton('Generate Report (PDF)')
        self.generate_pdf_button.clicked.connect(self.generate_report_pdf)
        button_layout.addWidget(self.generate_pdf_button)
        self.print_button = QPushButton('Print Report')
        self.print_button.clicked.connect(self.print_report)
        button_layout.addWidget(self.print_button)
        main_layout.addLayout(button_layout)
        rubric_layout = QHBoxLayout()
        self.manage_rubrics_button = QPushButton("Manage Rubrics")
        self.manage_rubrics_button.clicked.connect(self.manage_rubrics)
        rubric_layout.addWidget(self.manage_rubrics_button)
        self.run_analysis_button = QPushButton("Run Analysis")
        self.run_analysis_button.clicked.connect(self.run_rubric_analysis)
        rubric_layout.addWidget(self.run_analysis_button)
        self.rubric_list_widget = QListWidget()
        self.rubric_list_widget.setPlaceholderText("Available Rubrics")
        self.rubric_list_widget.setMaximumHeight(100)
        rubric_layout.addWidget(self.rubric_list_widget)
        main_layout.addLayout(rubric_layout)
        progress_layout = QHBoxLayout()
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 1)
        progress_layout.addWidget(self.progress_bar)
        self.cancel_button = QPushButton('Cancel Analysis')
        self.cancel_button.setEnabled(False)
        self.cancel_button.clicked.connect(self.cancel_analysis)
        progress_layout.addWidget(self.cancel_button)
        main_layout.addLayout(progress_layout)
        self.document_display_area = QTextEdit()
        self.document_display_area.setPlaceholderText("Drag and drop documents here, or use the 'Upload Document' button.")
        self.document_display_area.setReadOnly(True)
        self.document_display_area.setAcceptDrops(True)
        main_layout.addWidget(self.document_display_area)
        self.spacy_nlp_results_area = QTextEdit()
        self.spacy_nlp_results_area.setPlaceholderText("SpaCy NLP results (Tokens & Sentences) will appear here.")
        self.spacy_nlp_results_area.setReadOnly(True)
        main_layout.addWidget(self.spacy_nlp_results_area)
        self.clinical_ner_results_area = QTextEdit()
        self.clinical_ner_results_area.setPlaceholderText("Clinical/Biomedical NER results will appear here.")
        self.clinical_ner_results_area.setReadOnly(True)
        main_layout.addWidget(self.clinical_ner_results_area)
        self.spacy_ner_results_area = QTextEdit()
        self.spacy_ner_results_area.setPlaceholderText("SpaCy NER results will appear here.")
        self.spacy_ner_results_area.setReadOnly(True)
        main_layout.addWidget(self.spacy_ner_results_area)
        self.analysis_results_area = QTextEdit()
        self.analysis_results_area.setPlaceholderText("Rubric analysis results will appear here.")
        self.analysis_results_area.setReadOnly(True)
        main_layout.addWidget(self.analysis_results_area)
        self.central_widget.setLayout(main_layout)
        self.load_rubrics_to_main_list()

    def open_file_dialog(self):
        file_name, _ = QFileDialog.getOpenFileName(self, 'Select Document', '', 'Supported Files (*.pdf *.docx *.xlsx *.xls *.csv *.png *.jpg *.jpeg *.gif *.bmp *.tiff);;All Files (*.*)')
        if file_name:
            self.process_document(file_name)

    def _toggle_scrub_setting(self, checked: bool):
        self.scrub_before_display = checked
        if self._current_raw_text:
            shown = scrub_phi(self._current_raw_text) if self.scrub_before_display else self._current_raw_text
            self.document_display_area.setText(shown)
            self.status_bar.showMessage(f"Scrub before display set to {'ON' if checked else 'OFF'}.")

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            super().dragEnterEvent(event)

    def dropEvent(self, event: QDropEvent):
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                file_path = url.toLocalFile()
                self.process_document(file_path)
            event.acceptProposedAction()
        else:
            super().dropEvent(event)

    def _set_busy(self, busy: bool):
        if busy:
            self.progress_bar.setRange(0, 0)
            self.cancel_button.setEnabled(True)
        else:
            self.progress_bar.setRange(0, 1)
            self.cancel_button.setEnabled(False)

    def process_document(self, file_path):
        self.status_bar.showMessage(f"Processing: {os.path.basename(file_path)}")
        self.document_display_area.setText("Processing document in background...")
        self._current_raw_text = ""
        self._current_sentences_with_source = []
        self._current_entities_spacy = ""
        self._current_entities_transformer = ""
        self._set_busy(True)
        self._doc_thread = QThread()
        self._doc_worker = DocumentWorker(file_path)
        self._doc_worker.moveToThread(self._doc_thread)
        self._doc_thread.started.connect(self._doc_worker.run)
        self._doc_worker.finished.connect(self._handle_doc_finished)
        self._doc_worker.error.connect(self._handle_doc_error)
        self._doc_worker.progress.connect(lambda p: self.status_bar.showMessage(f"Document processing... {p}%"))
        self._doc_worker.finished.connect(self._doc_thread.quit)
        self._doc_worker.finished.connect(self._doc_worker.deleteLater)
        self._doc_thread.finished.connect(self._doc_thread.deleteLater)
        self._doc_thread.start()

    def cancel_analysis(self):
        canceled = False
        if hasattr(self, "_doc_worker") and self._doc_worker is not None:
            try:
                self._doc_worker.cancel()
                canceled = True
            except Exception:
                pass
        if hasattr(self, "_ner_worker") and self._ner_worker is not None:
            try:
                self._ner_worker.cancel()
                canceled = True
            except Exception:
                pass
        if hasattr(self, "_analysis_worker") and self._analysis_worker is not None:
            try:
                self._analysis_worker.cancel()
                canceled = True
            except Exception:
                pass
        if canceled:
            self.status_bar.showMessage("Cancel requested...")
        else:
            self.status_bar.showMessage("Nothing to cancel.")

    def _handle_doc_finished(self, sentences_with_source: List[Tuple[str, str]]):
        self._set_busy(False)
        self._current_sentences_with_source = sentences_with_source
        self._current_raw_text = "\n".join([text for text, source in sentences_with_source])
        display_text = self._current_raw_text
        if isinstance(display_text, str) and len(display_text) > 100000:
            display_text = display_text[:100000] + "\n...[truncated for display]"
        if self.scrub_before_display:
            display_text = scrub_phi(display_text)
        self.document_display_area.setText(display_text)
        self.status_bar.showMessage("Document processed.")
        if not self._current_raw_text or self._current_raw_text.startswith("Unsupported file type") or self._current_raw_text.startswith("Error processing file"):
            return
        self.process_spacy_basic_nlp(self._current_raw_text)
        self.process_spacy_ner(self._current_raw_text)
        self.process_biomedical_ner_async(self._current_raw_text)

    def _handle_doc_error(self, message: str):
        self._set_busy(False)
        self.document_display_area.setText(message)
        self.status_bar.showMessage("Document processing failed.")

    def process_spacy_basic_nlp(self, text):
        if not nlp or not text.strip():
            self.spacy_nlp_results_area.setText("SpaCy model not available (offline and not installed).")
            return
        doc = nlp(text)
        tokens = [token.text for token in doc]
        sentences = [sent.text for sent in doc.sents]
        nlp_output = (
            "--- Tokens ---\n" +
            "\n".join(tokens) +
            "\n\n--- Sentences ---\n" +
            "\n".join(sentences)
        )
        self.spacy_nlp_results_area.setText(nlp_output)
        self.status_bar.showMessage("SpaCy basic NLP complete.")

    def process_biomedical_ner_async(self, text: str):
        self.clinical_ner_results_area.setText("Running Clinical/Biomedical NER...")
        self.status_bar.showMessage("Running Clinical/Biomedical NER in background...")
        self._set_busy(True)
        self._ner_thread = QThread()
        self._ner_worker = NERWorker(text)
        self._ner_worker.moveToThread(self._ner_thread)
        self._ner_thread.started.connect(self._ner_worker.run)
        self._ner_worker.finished.connect(self._handle_ner_finished)
        self._ner_worker.error.connect(self._handle_ner_error)
        self._ner_worker.progress.connect(lambda p: self.status_bar.showMessage(f"NER... {p}%"))
        self._ner_worker.finished.connect(self._ner_thread.quit)
        self._ner_worker.finished.connect(self._ner_worker.deleteLater)
        self._ner_thread.finished.connect(self._ner_thread.deleteLater)
        self._ner_thread.start()

    def _handle_ner_finished(self, model_name: str, formatted_output: str):
        self._set_busy(False)
        self._current_entities_transformer = formatted_output
        self.clinical_ner_results_area.setText(f"[Model: {model_name}]\n{formatted_output}")
        self.status_bar.showMessage("Clinical/Biomedical NER processing complete.")

    def _handle_ner_error(self, message: str):
        self._set_busy(False)
        self.clinical_ner_results_area.setText(message)
        self.status_bar.showMessage("Clinical/Biomedical NER processing failed.")

    def process_spacy_ner(self, text):
        if not nlp or not text.strip():
            self.spacy_ner_results_area.setText("SpaCy model not available (offline and not installed).")
            return
        self.status_bar.showMessage("Running SpaCy NER...")
        try:
            doc = nlp(text)
            formatted_results = [f"Entity: {ent.text} | Type: {ent.label_}" for ent in doc.ents]
            self._current_entities_spacy = "\n".join(formatted_results) if formatted_results else "No entities detected by SpaCy."
            self.spacy_ner_results_area.setText(self._current_entities_spacy)
            self.status_bar.showMessage("SpaCy NER processing complete.")
        except Exception as e:
            self.spacy_ner_results_area.setText(f"Error during SpaCy NER: {e}")
            self.status_bar.showMessage("SpaCy NER processing failed.")

    def clear_document_display(self):
        self.document_display_area.clear()
        self.spacy_nlp_results_area.clear()
        self.clinical_ner_results_area.clear()
        self.spacy_ner_results_area.clear()
        self.analysis_results_area.clear()
        self._current_raw_text = ""
        self._current_sentences_with_source = []
        self._current_entities_spacy = ""
        self._current_entities_transformer = ""
        self.status_bar.showMessage('Display cleared.')

    def manage_rubrics(self):
        dialog = RubricManagerDialog(self)
        dialog.exec()
        self.load_rubrics_to_main_list()

    def load_rubrics_to_main_list(self):
        self.rubric_list_widget.clear()
        try:
            if not os.path.exists(DATABASE_PATH):
                return
            with sqlite3.connect(DATABASE_PATH) as conn:
                cur = conn.cursor()
                cur.execute("SELECT id, name FROM rubrics ORDER BY name ASC")
                for rubric_id, name in cur.fetchall():
                    item = QListWidgetItem(name)
                    item.setData(Qt.ItemDataRole.UserRole, rubric_id)
                    self.rubric_list_widget.addItem(item)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load rubrics into main list:\n{e}")

    def run_rubric_analysis(self):
        if not self._current_sentences_with_source:
            QMessageBox.warning(self, "Analysis Error", "Please upload a document to analyze first.")
            return
        selected_items = self.rubric_list_widget.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Analysis Error", "Please select a rubric from the list to run the analysis.")
            return
        try:
            rubric_id = selected_items[0].data(Qt.ItemDataRole.UserRole)
            with sqlite3.connect(DATABASE_PATH) as conn:
                cur = conn.cursor()
                cur.execute("SELECT content FROM rubrics WHERE id = ?", (rubric_id,))
                result = cur.fetchone()
            if not result:
                QMessageBox.critical(self, "Database Error", "Could not find the selected rubric in the database.")
                return
            rubric_content = result[0]
        except Exception as e:
            QMessageBox.critical(self, "Database Error", f"Failed to retrieve rubric content:\n{e}")
            return
        self.analysis_results_area.setText("Running rubric analysis...")
        self.status_bar.showMessage("Starting rubric analysis in background...")
        self._set_busy(True)
        self.progress_bar.setRange(0, 100)
        self._analysis_thread = QThread()
        self._analysis_worker = AnalysisWorker(self._current_sentences_with_source, rubric_content, OFFLINE_ONLY)
        self._analysis_worker.moveToThread(self._analysis_thread)
        self._analysis_thread.started.connect(self._analysis_worker.run)
        self._analysis_worker.finished.connect(self._handle_analysis_finished)
        self._analysis_worker.error.connect(self._handle_analysis_error)
        self._analysis_worker.progress.connect(lambda p: self.status_bar.showMessage(f"Rubric analysis... {p}%"))
        self._analysis_worker.progress.connect(self.progress_bar.setValue)
        self._analysis_worker.finished.connect(self._analysis_thread.quit)
        self._analysis_worker.finished.connect(self._analysis_worker.deleteLater)
        self._analysis_thread.finished.connect(self._analysis_thread.deleteLater)
        self._analysis_thread.start()

    def _handle_analysis_finished(self, report: str):
        self._set_busy(False)
        self.analysis_results_area.setText(report)
        self.status_bar.showMessage("Rubric analysis complete.")
        self.progress_bar.setRange(0, 1)

    def _handle_analysis_error(self, message: str):
        self._set_busy(False)
        self.analysis_results_area.setText(f"Analysis Failed:\n{message}")
        self.status_bar.showMessage("Rubric analysis failed.")
        self.progress_bar.setRange(0, 1)

    def _build_report_html(self) -> str:
        text_for_report = scrub_phi(self._current_raw_text or "")
        spacy_entities = scrub_phi(self._current_entities_spacy or "")
        hf_entities = scrub_phi(self._current_entities_transformer or "")
        html = f"""<html><head><meta charset=\"utf-8\"><style>body {{ font-family: Arial, sans-serif; }} h1 {{ font-size: 18pt; }} h2 {{ font-size: 14pt; margin-top: 12pt; }} pre {{ white-space: pre-wrap; font-family: Consolas, monospace; background: #f4f4f4; padding: 8px; }}</style></head><body><h1>Therapy Compliance Analysis Report</h1><p><b>Mode:</b> Offline | <b>PHI Scrubbing:</b> Enabled for export</p><h2>Extracted Text (scrubbed)</h2><pre>{text_for_report}</pre><h2>SpaCy Entities (scrubbed)</h2><h2>Clinical/Biomedical NER (scrubbed)</h2><pre>{hf_entities}</pre></body></html>"""
        return html

    def generate_report_pdf(self):
        if not self._current_raw_text:
            QMessageBox.information(self, "Generate Report", "No document data to export.")
            return
        save_path, _ = QFileDialog.getSaveFileName(self, 'Save Report as PDF', '', 'PDF Files (*.pdf)')
        if not save_path:
            return
        try:
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            if not save_path.lower().endswith(".pdf"):
                save_path += ".pdf"
            printer.setOutputFileName(save_path)
            from PyQt6.QtGui import QTextDocument
            doc = QTextDocument()
            doc.setHtml(self._build_report_html())
            doc.print(printer)
            QMessageBox.information(self, "Generate Report", f"PDF saved to:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Generate Report", f"Failed to create PDF:\n{e}")

    def print_report(self):
        if not self._current_raw_text:
            QMessageBox.information(self, "Print Report", "No document data to print.")
            return
        try:
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            dialog = QPrintDialog(printer, self)
            if dialog.exec() == QPrintDialog.DialogCode.Accepted:
                from PyQt6.QtGui import QTextDocument
                doc = QTextDocument()
                doc.setHtml(self._build_report_html())
                doc.print(printer)
                self.status_bar.showMessage("Report sent to printer.")
        except Exception as e:
            QMessageBox.critical(self, "Print Report", f"Failed to print:\n{e}")

    def show_paths(self):
        msg = f"App path:\n{os.path.abspath(__file__)}\n\nDatabase path:\n{os.path.abspath(DATABASE_PATH)}"
        QMessageBox.information(self, "Paths", msg)

    def show_quickstart(self):
        tips = (
            "Quickstart (Offline HIPAA Mode):\n"
            "1) Ensure SpaCy and NER models are pre-cached locally in this virtualenv.\n"
            "2) Windows OCR: Install Tesseract and set TESSERACT_EXE env var if needed.\n"
            "3) Click 'Upload Document' or drag & drop a file.\n"
            "4) Text shows (scrubbed), SpaCy results, then Clinical NER finishes.\n"
            "5) Use 'Generate Report (PDF)' or 'Print Report' for outputs.\n"
            "6) Use 'Cancel Analysis' to stop long runs.\n"
            "7) Tools → Verify Offline Readiness checks that everything is in place."
        )
        QMessageBox.information(self, "Quickstart", tips)

    def verify_offline_readiness(self):
        spacy_ok = nlp is not None
        spacy_msg = "OK" if spacy_ok else "Missing (install en_core_web_sm in this venv)."
        ner_ok = clinical_ner_pipeline is not None
        model_msg = _loaded_model_name if ner_ok else "Missing (pre-cache NER model in this venv)."
        tesseract_cmd = getattr(pytesseract.pytesseract, "tesseract_cmd", "")
        if tesseract_cmd and os.path.isfile(tesseract_cmd):
            tess_msg = f"OK ({tesseract_cmd})"
        else:
            tess_msg = "Unknown path (set TESSERACT_EXE env var or ensure in PATH)."
        try:
            import torch
            torch_msg = "OK"
        except Exception as e:
            torch_msg = f"Missing ({e})"
        msg = (
            f"Offline Readiness:\n\n"
            f"- SpaCy model: {spacy_msg}\n"
            f"- Transformers NER model: {model_msg}\n"
            f"- Tesseract: {tess_msg}\n"
            f"- PyTorch: {torch_msg}\n"
            f"- Offline mode: {'ENABLED' if OFFLINE_ONLY else 'Disabled'}"
        )
        QMessageBox.information(self, "Verify Offline Readiness", msg)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    initialize_database()
    main_win = MainApplicationWindow()
    main_win.show()
    sys.exit(app.exec())